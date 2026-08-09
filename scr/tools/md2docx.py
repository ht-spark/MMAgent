"""Markdown 转 DOCX 转换器 — 生成符合数学建模竞赛格式的 Word 文档。

功能：
  1. 解析 Markdown 文本（标题、段落、列表、表格、公式、图片引用）
  2. 生成格式化的 DOCX 文档：
     - 封面页（标题、摘要、关键词）
     - 自动编号的章节标题
     - 三线表格式
     - 居中图注
     - 居中公式（LaTeX → 文本近似）
     - 页码、页眉
  3. 嵌入实际图片文件（PNG）

依赖：python-docx
"""
from __future__ import annotations

import re
import shutil
import subprocess
import sys
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError as _e:
    # 可能安装了旧版 docx.py（Python 2 遗留单文件）而非 python-docx 包
    _msg = (
        f"无法导入 python-docx 包: {_e}\n"
        "可能原因：环境中存在旧版 docx.py 单文件（Python 2 遗留），"
        "遮蔽了正确的 python-docx 包。\n"
        "解决方法：pip install --force-reinstall python-docx"
    )
    raise ImportError(_msg) from _e

__all__ = [
    "markdown_to_docx",
    "convert_paper_md_to_docx",
    "find_pandoc",
    "pandoc_available",
    "pandoc_to_docx",
]


# ---------------------------------------------------------------------------
# Markdown 解析器
# ---------------------------------------------------------------------------

class MarkdownBlock:
    """Markdown 块级元素。"""
    HEADING = "heading"
    PARAGRAPH = "paragraph"
    TABLE = "table"
    FORMULA = "formula"
    IMAGE = "image"
    LIST = "list"
    HORIZONTAL_RULE = "hr"

    def __init__(self, block_type: str, content: str, level: int = 0, **kwargs):
        self.block_type = block_type
        self.content = content
        self.level = level
        self.extra = kwargs


def parse_markdown(md_text: str) -> list[MarkdownBlock]:
    """将 Markdown 文本解析为块级元素列表。"""
    lines = md_text.split("\n")
    blocks: list[MarkdownBlock] = []
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        # 空行跳过
        if not line.strip():
            i += 1
            continue

        # 标题
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            blocks.append(MarkdownBlock(MarkdownBlock.HEADING, title, level=level))
            i += 1
            continue

        # 水平线
        if re.match(r"^---+$", line.strip()):
            blocks.append(MarkdownBlock(MarkdownBlock.HORIZONTAL_RULE, ""))
            i += 1
            continue

        # 公式块 $$...$$
        if line.strip().startswith("$$"):
            formula_lines = [line.strip()]
            if not line.strip().endswith("$$") or line.strip() == "$$":
                i += 1
                while i < len(lines):
                    formula_lines.append(lines[i].rstrip())
                    if lines[i].strip().endswith("$$"):
                        break
                    i += 1
            formula_text = "\n".join(formula_lines)
            formula_text = formula_text.replace("$$", "").strip()
            blocks.append(MarkdownBlock(MarkdownBlock.FORMULA, formula_text))
            i += 1
            continue

        # 行内公式 \[...\]
        inline_formula = re.match(r"^\\\[([^\]]+)\\\]$", line.strip())
        if inline_formula:
            blocks.append(MarkdownBlock(MarkdownBlock.FORMULA, inline_formula.group(1).strip()))
            i += 1
            continue

        # 表格
        if line.startswith("|"):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            blocks.append(MarkdownBlock(MarkdownBlock.TABLE, "\n".join(table_lines)))
            continue

        # 图片引用 ![alt](path)
        img_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", line.strip())
        if img_match:
            alt = img_match.group(1)
            path = img_match.group(2)
            blocks.append(MarkdownBlock(MarkdownBlock.IMAGE, path, level=0, alt_text=alt))
            i += 1
            continue

        # 列表项
        list_match = re.match(r"^(\s*)([-*]|\d+\.)\s+(.+)$", line)
        if list_match:
            list_lines = [line]
            i += 1
            while i < len(lines):
                next_line = lines[i].rstrip()
                if re.match(r"^(\s*)([-*]|\d+\.)\s+", next_line):
                    list_lines.append(next_line)
                    i += 1
                elif next_line.strip() == "":
                    i += 1
                    # 检查下一行是否还是列表
                    if i < len(lines) and re.match(r"^(\s*)([-*]|\d+\.)\s+", lines[i].rstrip()):
                        continue
                    break
                else:
                    break
            blocks.append(MarkdownBlock(MarkdownBlock.LIST, "\n".join(list_lines)))
            continue

        # 普通段落（可能跨行直到空行或特殊行）
        para_lines = [line]
        i += 1
        while i < len(lines):
            next_line = lines[i].rstrip()
            if (not next_line.strip() or
                    next_line.startswith("#") or
                    next_line.startswith("|") or
                    next_line.startswith("$$") or
                    next_line.startswith("![") or
                    re.match(r"^---+$", next_line.strip()) or
                    re.match(r"^(\s*)([-*]|\d+\.)\s+", next_line)):
                break
            para_lines.append(next_line)
            i += 1
        blocks.append(MarkdownBlock(MarkdownBlock.PARAGRAPH, "\n".join(para_lines)))

    return blocks


# ---------------------------------------------------------------------------
# LaTeX → OMML 递归下降解析器
# ---------------------------------------------------------------------------

class _LatexParser:
    """LaTeX 公式递归下降解析器，输出 OMML XML 片段列表。

    支持：
      - 分数 \\frac{num}{den}
      - 求和/积分/乘积 \\sum_{a}^{b} \\int_{a}^{b} \\prod_{a}^{b}
      - 上下标 _{...} ^{...} 以及 _x ^x 单字符形式
      - 文本 \\text{...} \\mathrm{...} \\mathbf{...} \\mathcal{...} \\mathbb{...}
      - 重音 \\hat{x} \\bar{x} \\tilde{x} \\vec{x}
      - 希腊字母和数学运算符（自动替换为 Unicode）
      - 括号分组 { ... }
      - \\xrightarrow{...} \\xleftarrow{...}
      - \\tag{...} 公式编号
    """

    M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

    GREEK = {
        r"\alpha": "α", r"\beta": "β", r"\gamma": "γ", r"\delta": "δ",
        r"\epsilon": "ε", r"\varepsilon": "ε", r"\zeta": "ζ", r"\eta": "η",
        r"\theta": "θ", r"\vartheta": "ϑ", r"\lambda": "λ", r"\mu": "μ",
        r"\nu": "ν", r"\xi": "ξ", r"\pi": "π", r"\rho": "ρ",
        r"\sigma": "σ", r"\tau": "τ", r"\phi": "φ", r"\varphi": "φ",
        r"\chi": "χ", r"\psi": "ψ", r"\omega": "ω",
        r"\Gamma": "Γ", r"\Delta": "Δ", r"\Theta": "Θ", r"\Lambda": "Λ",
        r"\Sigma": "Σ", r"\Phi": "Φ", r"\Psi": "Ψ", r"\Omega": "Ω",
    }

    OPERATORS = {
        r"\leq": "≤", r"\geq": "≥", r"\neq": "≠", r"\ne": "≠",
        r"\times": "×", r"\div": "÷", r"\pm": "±", r"\mp": "∓",
        r"\infty": "∞", r"\partial": "∂",
        r"\forall": "∀", r"\exists": "∃", r"\nexists": "∄",
        r"\in": "∈", r"\notin": "∉", r"\ni": "∋",
        r"\cdot": "·", r"\ldots": "…", r"\cdots": "⋯", r"\vdots": "⋮", r"\ddots": "⋱",
        r"\to": "→", r"\rightarrow": "→", r"\leftarrow": "←", r"\leftrightarrow": "↔",
        r"\Rightarrow": "⇒", r"\Leftarrow": "⇐", r"\Leftrightarrow": "⇔",
        r"\cup": "∪", r"\cap": "∩", r"\subset": "⊂", r"\supset": "⊃",
        r"\subseteq": "⊆", r"\supseteq": "⊇",
        r"\approx": "≈", r"\equiv": "≡", r"\sim": "∼", r"\propto": "∝",
        r"\nabla": "∇", r"\angle": "∠", r"\perp": "⊥", r"\parallel": "∥",
        r"\land": "∧", r"\lor": "∨", r"\lnot": "¬", r"\neg": "¬",
        r"\oplus": "⊕", r"\ominus": "⊖", r"\otimes": "⊗", r"\odot": "⊙",
        r"\emptyset": "∅", r"\varnothing": "∅",
        r"\leqslant": "≤", r"\geqslant": "≥",
        r"\doteq": "≐", r"\models": "⊨", r"\vdash": "⊢", r"\dashv": "⊣",
    }

    # 文本类命令（提取 {} 内容作为普通文本）
    TEXT_CMDS = {r"\text", r"\mathrm", r"\mathbf", r"\mathcal",
                 r"\boldsymbol", r"\textbf", r"\textit", r"\operatorname"}

    # 重音命令
    ACCENT_MAP = {
        r"\hat": "\u0302",    # ̂
        r"\bar": "\u0304",    # ̄
        r"\tilde": "\u0303",  # ~
        r"\vec": "\u20D7",    # →
        r"\dot": "\u0307",    # ̇
        r"\ddot": "\u0308",   # ̈
    }

    # N元运算符（求和、积分、乘积）
    NARY_MAP = {
        r"\sum": "sum", r"\int": "int", r"\prod": "prod",
        r"\bigcup": "cup", r"\bigcap": "cap", r"\oint": "int",
    }

    # 函数名（直接输出为文本）
    FUNC_NAMES = {
        r"\min": "min", r"\max": "max", r"\ln": "ln", r"\log": "log",
        r"\exp": "exp", r"\sin": "sin", r"\cos": "cos", r"\tan": "tan",
        r"\arcsin": "arcsin", r"\arccos": "arccos", r"\arctan": "arctan",
        r"\sup": "sup", r"\inf": "inf", r"\arg": "arg", r"\deg": "deg",
        r"\det": "det", r"\dim": "dim", r"\gcd": "gcd", r"\lim": "lim",
        r"\top": "T",
    }

    def __init__(self, text: str):
        self.text = text
        self.pos = 0
        self.n = len(text)

    def parse(self) -> list[str]:
        """解析整个公式，返回 OMML 片段列表。"""
        parts: list[str] = []
        while self.pos < self.n:
            part = self._parse_atom()
            if part:
                parts.append(part)
        return parts

    def _peek(self, offset: int = 0) -> str:
        idx = self.pos + offset
        if idx < self.n:
            return self.text[idx]
        return ""

    def _match(self, s: str) -> bool:
        """尝试匹配字符串，成功则前进。"""
        if self.text[self.pos:self.pos + len(s)] == s:
            self.pos += len(s)
            return True
        return False

    def _read_command(self) -> str:
        """读取一个 \\ 命令（如 \\frac, \\alpha 等），返回含 \\ 的完整命令。"""
        start = self.pos
        self.pos += 1  # 跳过 \
        while self.pos < self.n and self.text[self.pos].isalpha():
            self.pos += 1
        # 处理单字符非字母命令（如 \!, \,, \;, \:）
        if self.pos == start + 1 and self.pos < self.n:
            ch = self.text[self.pos]
            if ch in "!,:;":
                self.pos += 1
        return self.text[start:self.pos]

    def _read_braced_group(self) -> str:
        """读取 {...} 分组内容（不含外层花括号），pos 停在 } 之后。"""
        # 跳过空白
        while self.pos < self.n and self.text[self.pos] in " \t":
            self.pos += 1
        if self.pos >= self.n or self.text[self.pos] != "{":
            return ""
        self.pos += 1  # 跳过 {
        depth = 1
        start = self.pos
        while self.pos < self.n and depth > 0:
            if self.text[self.pos] == "{":
                depth += 1
            elif self.text[self.pos] == "}":
                depth -= 1
                if depth == 0:
                    break
            self.pos += 1
        content = self.text[start:self.pos]
        if self.pos < self.n:
            self.pos += 1  # 跳过 }
        return content

    def _parse_atom(self) -> str:
        """解析单个原子，返回 OMML 片段。"""
        if self.pos >= self.n:
            return ""

        ch = self.text[self.pos]

        # 空白
        if ch in " \t\n":
            self.pos += 1
            return ""

        # 反斜杠命令
        if ch == "\\":
            return self._parse_script_chain(self._parse_command())

        # 下标 _
        if ch == "_":
            return self._parse_subscript()

        # 上标 ^
        if ch == "^":
            return self._parse_superscript()

        # 花括号分组
        if ch == "{":
            content = self._read_braced_group()
            sub_parser = _LatexParser(content)
            inner_parts = sub_parser.parse()
            return self._parse_script_chain("".join(inner_parts))

        # 普通字符
        self.pos += 1
        return self._parse_script_chain(self._make_run(ch))

    def _parse_script_chain(self, base: str) -> str:
        """Attach following _ and ^ markers to the parsed base atom."""
        if not base:
            return base

        while self.pos < self.n and self.text[self.pos] in "_^":
            sub_omml = ""
            sup_omml = ""

            while self.pos < self.n and self.text[self.pos] in "_^":
                marker = self.text[self.pos]
                self.pos += 1
                script = self._read_script_omml()
                if marker == "_":
                    sub_omml = script
                else:
                    sup_omml = script

            if sub_omml and sup_omml:
                base = (
                    f'<m:sSubSup xmlns:m="{self.M_NS}">'
                    f'<m:e>{base}</m:e>'
                    f'<m:sub>{sub_omml}</m:sub>'
                    f'<m:sup>{sup_omml}</m:sup>'
                    f'</m:sSubSup>'
                )
            elif sub_omml:
                base = (
                    f'<m:sSub xmlns:m="{self.M_NS}">'
                    f'<m:e>{base}</m:e>'
                    f'<m:sub>{sub_omml}</m:sub>'
                    f'</m:sSub>'
                )
            elif sup_omml:
                base = (
                    f'<m:sSup xmlns:m="{self.M_NS}">'
                    f'<m:e>{base}</m:e>'
                    f'<m:sup>{sup_omml}</m:sup>'
                    f'</m:sSup>'
                )

        return base

    def _read_script_omml(self) -> str:
        """Read the expression that follows _ or ^ as an OMML fragment."""
        while self.pos < self.n and self.text[self.pos] in " \t":
            self.pos += 1
        if self.pos >= self.n:
            return ""

        if self.text[self.pos] == "{":
            content = self._read_braced_group()
            sub_parser = _LatexParser(content)
            return "".join(sub_parser.parse())

        if self.text[self.pos] == "\\":
            return self._parse_command()

        ch = self.text[self.pos]
        self.pos += 1
        return self._make_run(ch)

    def _parse_command(self) -> str:
        """解析 \\ 命令。"""
        cmd = self._read_command()

        # 符号替换（希腊字母、运算符）
        if cmd in self.GREEK:
            return self._make_run(self.GREEK[cmd])
        if cmd in self.OPERATORS:
            return self._make_run(self.OPERATORS[cmd])

        # 文本类命令
        if cmd in self.TEXT_CMDS:
            content = self._read_braced_group()
            sub_parser = _LatexParser(content)
            inner = sub_parser.parse()
            return f'<m:r xmlns:m="{self.M_NS}"><m:rPr><m:sty m:val="p"/></m:rPr><m:t>{self._escape_xml("".join(self._extract_text(inner)))}</m:t></m:r>'

        # 分数 \frac{num}{den}
        if cmd == r"\frac":
            num_str = self._read_braced_group()
            den_str = self._read_braced_group()
            num_parser = _LatexParser(num_str)
            den_parser = _LatexParser(den_str)
            num_omml = "".join(num_parser.parse())
            den_omml = "".join(den_parser.parse())
            return (
                f'<m:f xmlns:m="{self.M_NS}">'
                f'<m:num>{num_omml}</m:num>'
                f'<m:den>{den_omml}</m:den>'
                f'</m:f>'
            )

        # \sqrt{...}
        if cmd == r"\sqrt":
            content = self._read_braced_group()
            sub_parser = _LatexParser(content)
            inner = "".join(sub_parser.parse())
            return f'<m:rad xmlns:m="{self.M_NS}"><m:deg></m:deg><m:e>{inner}</m:e></m:rad>'

        # \sqrt[n]{...}
        if cmd == r"\root" or (cmd == r"\sqrt" and self._peek() == "["):
            pass  # 简化处理

        # 重音命令
        if cmd in self.ACCENT_MAP:
            # 读取下一个原子作为被修饰的字符
            base = self._parse_atom()
            accent_char = self.ACCENT_MAP[cmd]
            return f'<m:acc xmlns:m="{self.M_NS}"><m:accPr><m:chr m:val="{accent_char}"/></m:accPr><m:e>{base}</m:e></m:acc>'

        # N元运算符（求和、积分等）
        if cmd in self.NARY_MAP:
            return self._parse_nary(cmd)

        # 函数名
        if cmd in self.FUNC_NAMES:
            return self._make_run(self.FUNC_NAMES[cmd])

        # \xrightarrow{} \xleftarrow{}
        if cmd in (r"\xrightarrow", r"\xleftarrow"):
            label = self._read_braced_group()
            arrow = "→" if "rightarrow" in cmd else "←"
            # 如果有标签（如 P, d），将标签作为上标显示
            label = label.strip()
            # 去掉间距命令
            label = label.replace(r"\;", "").replace(r"\,", "").replace(r"\!", "").strip()
            if label:
                label_omml = self._make_run(label)
                return (
                    f'<m:sSup xmlns:m="{self.M_NS}">'
                    f'<m:e>{self._make_run(arrow)}</m:e>'
                    f'<m:sup>{label_omml}</m:sup>'
                    f'</m:sSup>'
                )
            return self._make_run(arrow)

        # \tag{...}
        if cmd == r"\tag":
            tag_content = self._read_braced_group()
            return self._make_run(f"  ({tag_content})")

        # \quad \qquad
        if cmd == r"\quad":
            return self._make_run("  ")
        if cmd == r"\qquad":
            return self._make_run("    ")

        # 间距命令（不产生输出或产生细小间距）
        if cmd in (r"\!", r"\thinspace", r"\medspace", r"\thickspace"):
            return ""
        if cmd in (r"\,", r"\;", r"\:"):
            return self._make_run(" ")

        # \left \right （已预处理，但如果残留则跳过）
        if cmd in (r"\left", r"\right"):
            return ""

        # 括号尺寸命令（忽略尺寸，直接输出后续括号）
        if cmd in (r"\bigl", r"\bigr", r"\Bigl", r"\Bigr",
                    r"\big", r"\Big", r"\biggl", r"\biggr",
                    r"\Biggl", r"\Biggr", r"\biggm", r"\Bigm"):
            return ""

        # \stackrel{above}{below} — 将 above 作为 below 的上标
        if cmd == r"\stackrel":
            above = self._read_braced_group()
            below = self._parse_atom()
            above_parser = _LatexParser(above)
            above_omml = " ".join(above_parser.parse())
            return (
                f'<m:sSup xmlns:m="{self.M_NS}">'
                f'<m:e>{below}</m:e>'
                f'<m:sup>{above_omml}</m:sup>'
                f'</m:sSup>'
            )

        # \mathbb{1} 等特殊字符
        if cmd == r"\mathbb":
            content = self._read_braced_group()
            # 黑板粗体映射
            bb_map = {"1": "𝟙", "0": "𝟘", "R": "ℝ", "N": "ℕ", "Z": "ℤ",
                       "Q": "ℚ", "C": "ℂ", "P": "ℙ", "E": "𝔼", "H": "ℍ"}
            mapped = bb_map.get(content.strip(), content)
            return self._make_run(mapped)

        # 未知命令：静默忽略，避免输出原始命令名
        return ""

    def _parse_nary(self, cmd: str) -> str:
        """解析 N元运算符（求和、积分等），支持 _{...}^{...} 上下限。"""
        nary_val = self.NARY_MAP[cmd]
        lower = ""
        upper = ""

        # 读取下标
        if self.pos < self.n and self.text[self.pos] == "_":
            self.pos += 1
            if self.pos < self.n and self.text[self.pos] == "{":
                lower = self._read_braced_group()
            else:
                lower = self.text[self.pos] if self.pos < self.n else ""
                self.pos += 1

        # 读取上标
        if self.pos < self.n and self.text[self.pos] == "^":
            self.pos += 1
            if self.pos < self.n and self.text[self.pos] == "{":
                upper = self._read_braced_group()
            else:
                upper = self.text[self.pos] if self.pos < self.n else ""
                self.pos += 1

        while self.pos < self.n and self.text[self.pos] in " \t\n":
            self.pos += 1

        # 读取运算符后的表达式（一个原子）
        body = self._parse_atom()

        # 构建 OMML
        lower_omml = ""
        if lower:
            sub_parser = _LatexParser(lower)
            lower_omml = f'<m:sub>{" ".join(sub_parser.parse())}</m:sub>'

        upper_omml = ""
        if upper:
            sub_parser = _LatexParser(upper)
            upper_omml = f'<m:sup>{" ".join(sub_parser.parse())}</m:sup>'

        return (
            f'<m:nary xmlns:m="{self.M_NS}" m:val="{nary_val}">'
            f'{lower_omml}{upper_omml}'
            f'<m:e>{body}</m:e>'
            f'</m:nary>'
        )

    def _parse_subscript(self) -> str:
        """解析下标 _{...} 或 _x。"""
        self.pos += 1  # 跳过 _
        base = ""

        # 收集之前的 base（实际上是后续处理）
        if self.pos < self.n:
            if self.text[self.pos] == "{":
                content = self._read_braced_group()
                sub_parser = _LatexParser(content)
                sub_omml = " ".join(sub_parser.parse())
            else:
                ch = self.text[self.pos]
                self.pos += 1
                sub_omml = self._make_run(ch)

            return (
                f'<m:sSub xmlns:m="{self.M_NS}">'
                f'<m:e></m:e>'
                f'<m:sub>{sub_omml}</m:sub>'
                f'</m:sSub>'
            )
        return ""

    def _parse_superscript(self) -> str:
        """解析上标 ^{...} 或 ^x。"""
        self.pos += 1  # 跳过 ^
        if self.pos < self.n:
            if self.text[self.pos] == "{":
                content = self._read_braced_group()
                sub_parser = _LatexParser(content)
                sup_omml = " ".join(sub_parser.parse())
            else:
                ch = self.text[self.pos]
                self.pos += 1
                sup_omml = self._make_run(ch)

            return (
                f'<m:sSup xmlns:m="{self.M_NS}">'
                f'<m:e></m:e>'
                f'<m:sup>{sup_omml}</m:sup>'
                f'</m:sSup>'
            )
        return ""

    @staticmethod
    def _make_run(text: str) -> str:
        """生成 OMML 文本运行元素。"""
        escaped = _LatexParser._escape_xml(text)
        if not escaped:
            return ""
        return f'<m:r xmlns:m="{_LatexParser.M_NS}"><m:t>{escaped}</m:t></m:r>'

    @staticmethod
    def _escape_xml(text: str) -> str:
        """转义 XML 特殊字符。"""
        return (text
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;"))

    @staticmethod
    def _extract_text(omml_parts: list[str]) -> list[str]:
        """从 OMML 片段中粗略提取文本内容（用于 text 类命令）。"""
        texts: list[str] = []
        for part in omml_parts:
            # 提取 <m:t>...</m:t> 中的内容
            import re as _re
            matches = _re.findall(r'<m:t[^>]*>([^<]*)</m:t>', part)
            texts.extend(matches)
        return texts


# ---------------------------------------------------------------------------
# DOCX 生成器
# ---------------------------------------------------------------------------

class DocxBuilder:
    """构建 DOCX 文档。"""

    def __init__(self, output_path: str, base_dir: str = ""):
        self.doc = Document()
        self.output_path = output_path
        self.base_dir = base_dir
        self._figure_counter = 0
        self._table_counter = 0
        self._setup_styles()

    def _setup_styles(self):
        """设置文档默认样式。"""
        # 设置正文字体
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "宋体"
        font.size = Pt(12)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        # 设置段落间距
        pf = style.paragraph_format
        pf.space_after = Pt(6)
        pf.line_spacing = 1.5

        # 页面设置
        for section in self.doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)

    def build(self, blocks: list[MarkdownBlock], title: str = ""):
        """从解析的块构建文档。"""
        # 添加标题页
        if title:
            self._add_title_page(title)

        # 处理每个块
        for block in blocks:
            self._process_block(block)

        # 添加页码
        self._add_page_numbers()

    def _add_title_page(self, title: str):
        """添加标题页。"""
        # 大标题
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.name = "黑体"
        run.font.size = Pt(22)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_before = Pt(100)
        p.paragraph_format.space_after = Pt(30)

        self.doc.add_page_break()

    def _process_block(self, block: MarkdownBlock):
        """处理单个块。"""
        if block.block_type == MarkdownBlock.HEADING:
            self._add_heading(block.content, block.level)
        elif block.block_type == MarkdownBlock.PARAGRAPH:
            self._add_paragraph(block.content)
        elif block.block_type == MarkdownBlock.TABLE:
            self._add_table(block.content)
        elif block.block_type == MarkdownBlock.FORMULA:
            self._add_formula(block.content)
        elif block.block_type == MarkdownBlock.IMAGE:
            self._add_image(block.content, block.extra.get("alt_text", ""))
        elif block.block_type == MarkdownBlock.LIST:
            self._add_list(block.content)
        elif block.block_type == MarkdownBlock.HORIZONTAL_RULE:
            pass  # 忽略水平线

    def _add_heading(self, text: str, level: int):
        """添加标题。"""
        heading_levels = {
            1: ("黑体", 16, True),  # 一级标题
            2: ("黑体", 14, True),  # 二级标题
            3: ("黑体", 13, True),  # 三级标题
            4: ("黑体", 12, True),  # 四级标题
        }
        font_name, font_size, bold = heading_levels.get(level, ("黑体", 12, True))

        # 使用 Word 内置 heading 样式
        heading_style = f"Heading {min(level, 4)}"
        p = self.doc.add_paragraph(style=heading_style)
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

    def _add_paragraph(self, text: str):
        """添加段落，处理行内格式。"""
        # 检测图注/表注行（**图 N**：或 **表 N**：）
        stripped = text.strip()
        is_caption = False
        if stripped.startswith("**图 ") or stripped.startswith("**表 "):
            is_caption = True

        p = self.doc.add_paragraph()
        if is_caption:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.first_line_indent = Cm(0)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # 首行缩进2字符（约0.74cm）
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)

        # 处理行内格式：粗体、斜体、行内公式
        self._add_formatted_runs(p, text, caption_style=is_caption)

    def _add_formatted_runs(self, paragraph, text: str, caption_style: bool = False):
        """添加格式化的文本运行（处理 **粗体**、*斜体*、$行内公式$）。"""
        # 分割文本，处理 **bold** 和 $formula$
        parts = re.split(r"(\*\*[^*]+\*\*|\$[^$]+\$|`[^`]+`)", text)

        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
                if caption_style:
                    run.font.size = Pt(10.5)
            elif part.startswith("$") and part.endswith("$"):
                formula = self._clean_formula_text(part)
                if not self._append_omml_to_paragraph(paragraph, formula):
                    run = paragraph.add_run(self._latex_to_unicode(formula))
                    run.font.name = "Cambria Math"
                    run.font.size = Pt(12)
            elif part.startswith("`") and part.endswith("`"):
                run = paragraph.add_run(part[1:-1])
                run.font.name = "Consolas"
                run.font.size = Pt(10.5)
            else:
                paragraph.add_run(part)

    def _add_table(self, table_text: str):
        """添加三线表。"""
        lines = [l.strip() for l in table_text.split("\n") if l.strip().startswith("|")]
        if len(lines) < 2:
            return

        # 解析表格行
        rows_data: list[list[str]] = []
        for line in lines:
            # 跳过分隔行 |---|---|
            if re.match(r"^\|[\s\-:|]+\|$", line):
                continue
            cells = [c.strip() for c in line.split("|")[1:-1]]
            rows_data.append(cells)

        if not rows_data:
            return

        n_cols = max(len(r) for r in rows_data)
        # 补齐短行
        for r in rows_data:
            while len(r) < n_cols:
                r.append("")

        # 创建表格
        table = self.doc.add_table(rows=len(rows_data), cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        # 计算各列最大宽度（用于自适应）
        col_max_len = [0] * n_cols
        for r in rows_data:
            for j, cell in enumerate(r):
                col_max_len[j] = max(col_max_len[j], len(cell))

        # 设置列宽（基于内容长度，但有上下限）
        total_len = sum(col_max_len)
        for j in range(n_cols):
            if total_len > 0:
                ratio = col_max_len[j] / total_len
                width_cm = max(2.0, min(8.0, ratio * 15.0))
                for row in table.rows:
                    row.cells[j].width = Cm(width_cm)

        # 填充数据
        for i, row_data in enumerate(rows_data):
            for j, cell_text in enumerate(row_data):
                cell = table.cell(i, j)
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 处理加粗
                clean_text = cell_text
                bold = False
                if clean_text.startswith("**") and clean_text.endswith("**"):
                    clean_text = clean_text[2:-2]
                    bold = True

                run = p.add_run(clean_text)
                run.font.name = "宋体"
                run.font.size = Pt(10.5)

                # 表头加粗
                if i == 0 or bold:
                    run.bold = True

                # 设置单元格内边距
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcMar = parse_xml(f'''<w:tcMar {nsdecls("w")}>
                    <w:top w:w="50" w:type="dxa"/>
                    <w:bottom w:w="50" w:type="dxa"/>
                    <w:left w:w="100" w:type="dxa"/>
                    <w:right w:w="100" w:type="dxa"/>
                </w:tcMar>''')
                tcPr.append(tcMar)

                # 设置最小行高
                tr = table.rows[i]
                trPr = tr._tr.get_or_add_trPr()
                trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="340" w:hRule="atLeast"/>')
                trPr.append(trHeight)

        # 三线表样式：顶部、表头下方、底部有线，其余无线
        self._format_three_line_table(table)

        self._table_counter += 1
        self.doc.add_paragraph()

    def _format_three_line_table(self, table):
        """设置三线表样式。"""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')

        # 移除默认边框
        existing_borders = tblPr.find(qn("w:tblBorders"))
        if existing_borders is not None:
            tblPr.remove(existing_borders)

        # 设置三线表边框（1.5磅顶线、1.5磅底线、0.75磅表头线）
        borders_xml = f'''<w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="18" w:space="0" w:color="000000"/>
            <w:bottom w:val="single" w:sz="18" w:space="0" w:color="000000"/>
            <w:left w:val="nil"/>
            <w:right w:val="nil"/>
            <w:insideH w:val="nil"/>
            <w:insideV w:val="nil"/>
        </w:tblBorders>'''
        tblPr.append(parse_xml(borders_xml))

        # 表头行底部添加边框
        if len(table.rows) > 1:
            first_row = table.rows[0]
            for cell in first_row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = parse_xml(f'''<w:tcBorders {nsdecls("w")}>
                    <w:bottom w:val="single" w:sz="9" w:space="0" w:color="000000"/>
                </w:tcBorders>''')
                tcPr.append(tcBorders)

    def _add_formula(self, formula_text: str):
        """添加居中公式段落（MathType 风格，使用 OMML 渲染）。

        将 LaTeX 公式转换为 OMML (Office Math Markup Language)，
        在 Word 中以原生数学公式格式显示，效果接近 MathType。
        """
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        formula_text = self._clean_formula_text(formula_text)

        if self._append_omml_to_paragraph(p, formula_text):
            return

        # 回退：使用 Unicode 近似
        readable = self._latex_to_unicode(formula_text)
        run = p.add_run(readable)
        run.font.name = "Cambria Math"
        run.font.size = Pt(12)
        run.italic = False

    def _append_omml_to_paragraph(self, paragraph, formula_text: str) -> bool:
        """Append a native Word math node to a paragraph."""
        omml_xml = self._latex_to_omml(formula_text)
        if not omml_xml:
            return False
        try:
            paragraph._p.append(parse_xml(omml_xml))
            return True
        except Exception:
            return False

    @staticmethod
    def _clean_formula_text(formula_text: str) -> str:
        """Remove Markdown math delimiters before LaTeX parsing."""
        formula = formula_text.strip()
        if formula.startswith("$$") and formula.endswith("$$"):
            formula = formula[2:-2].strip()
        elif formula.startswith("$") and formula.endswith("$"):
            formula = formula[1:-1].strip()
        if formula.startswith(r"\[") and formula.endswith(r"\]"):
            formula = formula[2:-2].strip()
        return re.sub(r"\s+", " ", formula).strip()

    # ------------------------------------------------------------------
    # LaTeX → OMML 递归下降解析器
    # ------------------------------------------------------------------

    # OMML 命名空间
    _M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

    # 符号映射表
    _GREEK = {
        r"\alpha": "α", r"\beta": "β", r"\gamma": "γ", r"\delta": "δ",
        r"\epsilon": "ε", r"\varepsilon": "ε", r"\zeta": "ζ", r"\eta": "η",
        r"\theta": "θ", r"\vartheta": "ϑ", r"\lambda": "λ", r"\mu": "μ",
        r"\nu": "ν", r"\xi": "ξ", r"\pi": "π", r"\rho": "ρ",
        r"\sigma": "σ", r"\tau": "τ", r"\phi": "φ", r"\varphi": "φ",
        r"\chi": "χ", r"\psi": "ψ", r"\omega": "ω",
        r"\Gamma": "Γ", r"\Delta": "Δ", r"\Theta": "Θ", r"\Lambda": "Λ",
        r"\Sigma": "Σ", r"\Phi": "Φ", r"\Psi": "Ψ", r"\Omega": "Ω",
    }

    _OPERATORS = {
        r"\leq": "≤", r"\geq": "≥", r"\neq": "≠", r"\ne": "≠",
        r"\times": "×", r"\div": "÷", r"\pm": "±", r"\mp": "∓",
        r"\infty": "∞", r"\partial": "∂",
        r"\forall": "∀", r"\exists": "∃", r"\nexists": "∄",
        r"\in": "∈", r"\notin": "∉", r"\ni": "∋",
        r"\cdot": "·", r"\ldots": "…", r"\cdots": "⋯", r"\vdots": "⋮", r"\ddots": "⋱",
        r"\to": "→", r"\rightarrow": "→", r"\leftarrow": "←", r"\leftrightarrow": "↔",
        r"\Rightarrow": "⇒", r"\Leftarrow": "⇐", r"\Leftrightarrow": "⇔",
        r"\cup": "∪", r"\cap": "∩", r"\subset": "⊂", r"\supset": "⊃",
        r"\subseteq": "⊆", r"\supseteq": "⊇",
        r"\approx": "≈", r"\equiv": "≡", r"\sim": "∼", r"\propto": "∝",
        r"\nabla": "∇", r"\angle": "∠", r"\perp": "⊥", r"\parallel": "∥",
    }

    _SYMBOL_MAP = {**_GREEK, **_OPERATORS}

    def _latex_to_omml(self, latex: str) -> str:
        """将 LaTeX 公式转换为 OMML XML（Office MathML）。

        使用递归下降解析器正确处理嵌套结构：
        分数 \\frac{}{}、求和 \\sum_{}^{}、上下标 _{} ^{}、
        文本 \\text{}、重音 \\hat{} \\bar{} 等。
        """
        M_NS = self._M_NS

        # 预处理：移除 \left \right \quad \qquad
        text = latex
        text = text.replace(r"\left", "").replace(r"\right", "")
        text = text.replace(r"\quad", "  ").replace(r"\qquad", "    ")

        # 解析为 OMML 片段列表
        parser = _LatexParser(text)
        omml_parts = parser.parse()

        if not omml_parts:
            return ""

        # 包装在 oMath 元素中
        inner = "".join(omml_parts)
        return f'<m:oMath xmlns:m="{M_NS}">{inner}</m:oMath>'

    def _latex_to_unicode(self, latex: str) -> str:
        """将 LaTeX 公式转换为 Unicode 近似文本（回退方案）。"""
        text = latex
        # 移除命令
        text = re.sub(r"\\text\{([^}]*)\}", r"\1", text)
        text = re.sub(r"\\mathrm\{([^}]*)\}", r"\1", text)
        text = re.sub(r"\\mathbf\{([^}]*)\}", r"\1", text)
        text = re.sub(r"\\mathcal\{([^}]*)\}", r"\1", text)
        text = re.sub(r"\\boldsymbol\{([^}]*)\}", r"\1", text)
        text = re.sub(r"\\mathbb\{([^}]*)\}", r"\1", text)
        text = text.replace(r"\left", "").replace(r"\right", "")
        text = text.replace(r"\quad", "  ").replace(r"\qquad", "    ")
        # 替换符号
        for latex_sym, unicode_sym in self._SYMBOL_MAP.items():
            text = text.replace(latex_sym, unicode_sym)
        # 简单命令
        text = text.replace(r"\sum", "∑").replace(r"\int", "∫")
        text = text.replace(r"\prod", "∏").replace(r"\min", "min")
        text = text.replace(r"\max", "max").replace(r"\ln", "ln")
        text = text.replace(r"\log", "log").replace(r"\exp", "exp")
        text = text.replace(r"\frac", "/")
        # 移除剩余的 LaTeX 命令
        text = re.sub(r"\\[a-zA-Z]+", "", text)
        # 清理括号
        text = text.replace("{", "(").replace("}", ")")
        return text.strip()

    def _add_image(self, image_path: str, alt_text: str = ""):
        """添加图片（不自动添加图注，由 Markdown 正文中的图注段落处理）。"""
        # 尝试解析路径
        full_path = image_path
        if not Path(full_path).exists() and self.base_dir:
            full_path = str(Path(self.base_dir) / image_path)

        if not Path(full_path).exists():
            # 图片不存在，添加占位文本
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"[图片缺失: {image_path}]")
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            return

        try:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(full_path, width=Inches(5.5))
        except Exception as e:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"[图片加载失败: {image_path} ({e})]")
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)

    def _add_list(self, list_text: str):
        """添加列表。"""
        lines = list_text.split("\n")
        for line in lines:
            match = re.match(r"^(\s*)([-*]|\d+\.)\s+(.+)$", line)
            if match:
                indent = len(match.group(1))
                marker = match.group(2)
                content = match.group(3)

                is_ordered = re.match(r"\d+\.", marker)
                list_style = "List Number" if is_ordered else "List Bullet"

                p = self.doc.add_paragraph(style=list_style)
                self._add_formatted_runs(p, content)

                # 设置缩进
                if indent > 0:
                    p.paragraph_format.left_indent = Cm(indent * 0.5 + 0.75)

    def _add_page_numbers(self):
        """添加页码。"""
        for section in self.doc.sections:
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加页码字段
            run = footer_para.add_run()
            fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
            run._r.append(fldChar1)

            run2 = footer_para.add_run()
            instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
            run2._r.append(instrText)

            run3 = footer_para.add_run()
            fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
            run3._r.append(fldChar2)

    def save(self):
        """保存文档。"""
        self.doc.save(self.output_path)
        print(f"[md2docx] DOCX 已保存: {self.output_path}")


# ---------------------------------------------------------------------------
# 公开接口
# ---------------------------------------------------------------------------


def _project_tools_dir() -> Path:
    """返回项目根目录下的 tools/ 目录。"""
    return Path(__file__).resolve().parent.parent.parent / "tools"


def find_pandoc() -> str | None:
    """查找系统或项目内可用的 Pandoc 可执行文件。"""
    exe_name = "pandoc.exe" if sys.platform == "win32" else "pandoc"
    found = shutil.which(exe_name)
    if found:
        return found

    for candidate in (
        _project_tools_dir() / "pandoc" / exe_name,
        _project_tools_dir() / exe_name,
    ):
        if candidate.exists():
            return str(candidate)

    if sys.platform == "win32":
        candidate = Path.home() / "AppData" / "Local" / "Pandoc" / exe_name
        if candidate.exists():
            return str(candidate)
    return None


def pandoc_available() -> bool:
    """Pandoc 是否可用。"""
    return find_pandoc() is not None


def _post_process_pandoc_docx(docx_path: str | Path) -> None:
    """后处理 Pandoc 生成的 DOCX，应用三线表和中文论文格式。

    Pandoc 生成的 DOCX 公式为原生 OMML（可编辑），但表格无三线表样式、
    字体为西文默认、缺少页码和中文排版规范。此函数在 Pandoc 转换后
    对 DOCX 进行格式修正。

    Args:
        docx_path: Pandoc 生成的 DOCX 文件路径。
    """
    doc = Document(str(docx_path))

    # -- 1. 文档默认样式 --
    _apply_document_styles(doc)

    # -- 2. 页面边距 --
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # -- 3. 表格三线表格式 --
    for table in doc.tables:
        _apply_three_line_table_format(table)

    # -- 4. 页码 --
    _add_page_numbers_to_doc(doc)

    doc.save(str(docx_path))
    print(f"[md2docx] 后处理完成（三线表+样式）: {docx_path}", flush=True)


def _apply_document_styles(doc: Document) -> None:
    """设置文档默认样式和标题样式。"""
    # Normal 样式：宋体 12pt，1.5 倍行距
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5

    # 标题样式：黑体
    heading_configs = {
        "Heading 1": (16, True),
        "Heading 2": (14, True),
        "Heading 3": (13, True),
        "Heading 4": (12, True),
    }
    for style_name, (font_size, bold) in heading_configs.items():
        try:
            hs = doc.styles[style_name]
        except KeyError:
            continue
        hs.font.name = "Times New Roman"
        hs.font.size = Pt(font_size)
        hs.font.bold = bold
        hs.font.color.rgb = RGBColor(0, 0, 0)
        rPr = hs.element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), "黑体")
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)


def _apply_three_line_table_format(table) -> None:
    """对单个表格应用三线表格式。

    - 顶部 1.5pt 粗线
    - 底部 1.5pt 粗线
    - 表头行底部 0.75pt 细线
    - 其余无边框
    - 列宽按内容自适应
    - 单元格宋体 10.5pt 居中，表头加粗
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)

    # 移除已有边框和样式
    for child_name in ("w:tblBorders", "w:tblStyle"):
        existing = tblPr.find(qn(child_name))
        if existing is not None:
            tblPr.remove(existing)

    # 设置三线表边框
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '<w:left w:val="nil"/>'
        '<w:right w:val="nil"/>'
        '<w:insideH w:val="nil"/>'
        '<w:insideV w:val="nil"/>'
        '</w:tblBorders>'
    )
    tblPr.append(parse_xml(borders_xml))

    # 设置表格宽度为页面可用宽度
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = parse_xml(f'<w:tblW {nsdecls("w")}/>')
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "8785")  # 约 15.5cm (页面宽度减去边距)
    tblW.set(qn("w:type"), "dxa")

    # 设置固定布局（使列宽设置生效）
    existing_layout = tblPr.find(qn("w:tblLayout"))
    if existing_layout is not None:
        tblPr.remove(existing_layout)
    tblLayout = parse_xml(
        f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>'
    )
    tblPr.append(tblLayout)

    # 计算列宽（按内容长度比例分配）
    n_cols = len(table.columns)
    n_rows = len(table.rows)
    if n_cols == 0 or n_rows == 0:
        return

    col_max_len = [0] * n_cols
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            text_len = max(len(line) for line in cell.text.split("\n")) if cell.text else 1
            # 中文字符算2个宽度
            weighted_len = sum(2 if ord(c) > 127 else 1 for c in cell.text)
            col_max_len[j] = max(col_max_len[j], weighted_len, text_len * 2)

    total_len = sum(col_max_len) or 1
    total_width_dxa = 8785  # 总宽度 (dxa)
    col_widths = []
    for j in range(n_cols):
        ratio = col_max_len[j] / total_len
        width = max(1200, min(4000, int(total_width_dxa * ratio)))
        col_widths.append(width)

    # 微调使总宽度匹配
    diff = total_width_dxa - sum(col_widths)
    if diff != 0:
        col_widths[-1] += diff

    # 更新 tblGrid 中的 gridCol 元素（决定实际列宽）
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is not None:
        gridCols = tblGrid.findall(qn("w:gridCol"))
        for j, gridCol in enumerate(gridCols):
            if j < len(col_widths):
                gridCol.set(qn("w:w"), str(col_widths[j]))

    # 设置列宽和单元格格式
    for i, row in enumerate(table.rows):
        # 设置行高
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = parse_xml(
            f'<w:trHeight {nsdecls("w")} w:val="340" w:hRule="atLeast"/>'
        )
        trPr.append(trHeight)

        for j, cell in enumerate(row.cells):
            # 设置列宽（同时更新 tcW）
            if j < len(col_widths):
                tc = cell._tc
                tcPr2 = tc.get_or_add_tcPr()
                existing_tcW = tcPr2.find(qn("w:tcW"))
                if existing_tcW is not None:
                    tcPr2.remove(existing_tcW)
                tcW_xml = parse_xml(
                    f'<w:tcW {nsdecls("w")} w:w="{col_widths[j]}" w:type="dxa"/>'
                )
                tcPr2.insert(0, tcW_xml)

            # 设置单元格内边距
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            # 移除已有边距设置
            existing_mar = tcPr.find(qn("w:tcMar"))
            if existing_mar is not None:
                tcPr.remove(existing_mar)
            tcMar = parse_xml(
                f'<w:tcMar {nsdecls("w")}>'
                '<w:top w:w="50" w:type="dxa"/>'
                '<w:bottom w:w="50" w:type="dxa"/>'
                '<w:left w:w="100" w:type="dxa"/>'
                '<w:right w:w="100" w:type="dxa"/>'
                '</w:tcMar>'
            )
            tcPr.append(tcMar)

            # 设置单元格段落格式
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = Cm(0)
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)

                # 设置 run 字体
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10.5)
                    run.font.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    # 表头行加粗
                    if i == 0:
                        run.bold = True

            # 表头行底部边框（第二行开始才需要）
            if i == 0 and n_rows > 1:
                existing_borders = tcPr.find(qn("w:tcBorders"))
                if existing_borders is not None:
                    tcPr.remove(existing_borders)
                tcBorders = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>'
                    '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
                    '</w:tcBorders>'
                )
                tcPr.append(tcBorders)


def _add_page_numbers_to_doc(doc: Document) -> None:
    """为文档添加页码（页脚居中）。"""
    for section in doc.sections:
        footer = section.footer
        footer_para = (
            footer.paragraphs[0]
            if footer.paragraphs
            else footer.add_paragraph()
        )
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 清除已有内容
        for run in footer_para.runs:
            run.text = ""

        # 添加页码字段
        run1 = footer_para.add_run()
        fldChar1 = parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
        )
        run1._r.append(fldChar1)

        run2 = footer_para.add_run()
        instrText = parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
        )
        run2._r.append(instrText)

        run3 = footer_para.add_run()
        fldChar2 = parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
        )
        run3._r.append(fldChar2)


def pandoc_to_docx(
    md_path: str | Path,
    output_path: str | Path,
    resource_path: str | Path | None = None,
    reference_doc: str | Path | None = None,
    extra_args: list[str] | None = None,
    post_process: bool = True,
) -> str:
    """调用 Pandoc 将 Markdown 转为 DOCX（含可编辑 Word 公式）。

    Args:
        md_path: Markdown 文件路径。
        output_path: 输出 DOCX 文件路径。
        resource_path: 资源（图片）基准目录。
        reference_doc: 参考文档（用于样式模板）。
        extra_args: 附加 Pandoc 命令行参数。
        post_process: 是否在后处理中应用三线表和中文论文格式。

    Returns:
        输出 DOCX 文件路径。
    """
    pandoc = find_pandoc()
    if pandoc is None:
        raise FileNotFoundError("系统未检测到 Pandoc")

    md_path = Path(md_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    resource_path = Path(resource_path) if resource_path else md_path.parent

    cmd = [
        pandoc,
        str(md_path),
        "-o", str(output_path),
        "--from", "markdown+tex_math_dollars+raw_tex",
        "--resource-path", str(resource_path),
    ]
    if reference_doc is not None:
        cmd += ["--reference-doc", str(reference_doc)]
    if extra_args:
        cmd += extra_args

    proc = subprocess.run(
        cmd, capture_output=True, text=True, encoding="utf-8", errors="replace"
    )
    if proc.returncode != 0:
        raise subprocess.CalledProcessError(
            proc.returncode, cmd, output=proc.stdout, stderr=proc.stderr[-2000:]
        )

    # 后处理：应用三线表和中文论文格式
    if post_process:
        try:
            _post_process_pandoc_docx(output_path)
        except Exception as exc:  # noqa: BLE001
            print(f"[md2docx] 后处理失败（不影响公式）: {exc}", flush=True)

    return str(output_path)


def markdown_to_docx(
    md_text: str,
    output_path: str,
    title: str = "",
    base_dir: str = "",
) -> str:
    """将 Markdown 文本转换为 DOCX 文档。

    Args:
        md_text: Markdown 文本。
        output_path: 输出 DOCX 文件路径。
        title: 文档标题（用于封面页）。
        base_dir: 图片基准目录（用于解析相对路径）。

    Returns:
        输出文件路径。
    """
    blocks = parse_markdown(md_text)
    builder = DocxBuilder(output_path, base_dir=base_dir)
    builder.build(blocks, title=title)
    builder.save()
    return output_path


def _preprocess_markdown_math(md_text: str) -> str:
    """预处理 Markdown 中的数学公式格式，提高 Pandoc 渲染成功率。

    修复以下问题：
      1. 行内公式 ``$ formula $`` → ``$formula$``
         Pandoc 的 tex_math_dollars 扩展要求 ``$`` 紧贴非空格字符，
         否则不识别为数学公式。生成的 Markdown 中行内公式常有额外空格。
      2. 段落内嵌入的 ``$$ ... $$`` 独立成行
         确保 Pandoc 将其正确识别为 display math 块。

    Args:
        md_text: 原始 Markdown 文本。

    Returns:
        预处理后的 Markdown 文本。
    """
    # 1. 修复行内公式：$ formula $ → $formula$
    #    (?<!\$) 避免 $$ 的第二个 $ 被匹配
    #    (?!\$)  避免 $$ 的第一个 $ 被匹配
    md_text = re.sub(
        r'(?<!\$)\$\s+([^\$\n]+?)\s+\$(?!\$)',
        r'$\1$',
        md_text,
    )

    # 2. 将行内嵌入的 $$ ... $$ 独立成行
    #    匹配 "文字$$formula$$文字" → "文字\n$$formula$$\n文字"
    #    只处理 $$ 前后有非换行字符的情况
    md_text = re.sub(
        r'(?<=\S)\$\$([^\$\n]+)\$\$',
        r'\n$$\1$$\n',
        md_text,
    )
    md_text = re.sub(
        r'\$\$([^\$\n]+)\$\$(?=\S)',
        r'$$\1$$\n',
        md_text,
    )

    # 3. 替换 \tag{N} → \qquad (N)
    #    Pandoc 的 tex_math_dollars 不支持 \tag{} 命令，
    #    会导致包含 \tag{} 的公式无法渲染。
    #    替换为 \qquad (N) 可在公式内显示编号。
    md_text = re.sub(
        r'\\tag\{([^}]+)\}',
        r'\\qquad (\1)',
        md_text,
    )

    return md_text


def convert_paper_md_to_docx(
    md_file_path: str,
    output_dir: str | None = None,
) -> str:
    """将报告 Markdown 文件转换为 DOCX。

    优先使用 Pandoc，以保留可编辑的 Word 原生公式；若未检测到 Pandoc
    或 Pandoc 转换失败，则明确记录日志并回退到 Python-docx 实现。

    Args:
        md_file_path: Markdown 文件路径。
        output_dir: 输出目录（默认与 md 文件同目录）。

    Returns:
        输出 DOCX 文件路径。
    """
    md_path = Path(md_file_path)
    if not md_path.exists():
        raise FileNotFoundError(f"Markdown 文件不存在: {md_file_path}")

    if output_dir is None:
        output_dir = str(md_path.parent)
    output_path = str(Path(output_dir) / f"{md_path.stem}.docx")

    if pandoc_available():
        try:
            print("[md2docx] 检测到 Pandoc，使用 Pandoc 转换 DOCX...", flush=True)
            # 预处理 Markdown：修复行内公式格式
            md_text = md_path.read_text(encoding="utf-8")
            md_text = _preprocess_markdown_math(md_text)
            # 写入临时文件（同目录，确保图片路径可解析）
            temp_md = md_path.parent / f".{md_path.stem}_pandoc_tmp.md"
            temp_md.write_text(md_text, encoding="utf-8")
            try:
                result = pandoc_to_docx(
                    temp_md, output_path, resource_path=md_path.parent
                )
            finally:
                temp_md.unlink(missing_ok=True)
            print(f"[md2docx] Pandoc 转换完成: {result}", flush=True)
            return result
        except Exception as exc:  # noqa: BLE001
            print(f"[md2docx] Pandoc 转换失败，将回退 Python-docx: {exc}", flush=True)
    else:
        print("[md2docx] 未检测到 Pandoc，使用 Python-docx 回退转换...", flush=True)

    md_text = md_path.read_text(encoding="utf-8")

    # 从 Markdown 提取标题
    title = ""
    title_match = re.match(r"^#\s+(.+)$", md_text, re.MULTILINE)
    if title_match:
        title = title_match.group(1).strip()
        # 移除标题行，避免重复
        md_text = md_text.replace(title_match.group(0), "", 1)

    # base_dir 用于解析图片路径
    base_dir = str(md_path.parent)

    result = markdown_to_docx(
        md_text,
        output_path,
        title=title,
        base_dir=base_dir,
    )
    print(f"[md2docx] Python-docx 转换完成: {result}", flush=True)
    return result
